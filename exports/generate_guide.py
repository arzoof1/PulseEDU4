#!/usr/bin/env python3
"""Generate the PulseEDU demo guide as both .docx and .pdf from one content model."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ---------------------------------------------------------------------------
# Brand
# ---------------------------------------------------------------------------
PRIMARY = "1F3A8A"      # deep indigo
ACCENT = "0E7C86"       # teal
INK = "1F2937"          # near-black
MUTE = "6B7280"         # gray
BOXBG = "F1F4FB"        # light box fill
BOXBORDER = "B9C4E6"

def hx(s):  # "1F3A8A" -> reportlab Color
    return colors.HexColor("#" + s)

def rgb(s):  # "1F3A8A" -> docx RGBColor
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))

PASSWORD = "PulseDemo26"
SITE = "PulseKinetics.us"
SCHOOL = "D. S. Parrott Middle School"

LOGINS_ELA = [
    ("Kathleen Taylor", "ELA \u2013 Grade 6", "kathleen.taylor@pulsedemo.com"),
    ("Donna Sanchez", "ELA \u2013 Grade 6", "donna.sanchez@pulsedemo.com"),
    ("Shirley Rodriguez", "ELA \u2013 Grade 7", "shirley.rodriguez@pulsedemo.com"),
    ("Pamela Martin", "ELA \u2013 Grade 7", "pamela.martin@pulsedemo.com"),
    ("Jennifer Harris", "ELA \u2013 Grade 8", "jennifer.harris@pulsedemo.com"),
    ("Steven Garcia", "ELA \u2013 Grade 8", "steven.garcia@pulsedemo.com"),
]
LOGINS_MATH = [
    ("Brian Martinez", "Math \u2013 Grade 6", "brian.martinez@pulsedemo.com"),
    ("Kathleen Sanchez", "Math \u2013 Grade 6", "kathleen.sanchez@pulsedemo.com"),
    ("Amy Brown", "Math \u2013 Grade 7", "amy.brown@pulsedemo.com"),
    ("Carol Garcia", "Math \u2013 Grade 7", "carol.garcia@pulsedemo.com"),
    ("Christine Green", "Math \u2013 Grade 8", "christine.green@pulsedemo.com"),
    ("Larry Martin", "Math \u2013 Grade 8", "larry.martin@pulsedemo.com"),
]

# Each feature: (title, overview, where, try_this, [bullets], placeholder_caption)
FEATURES = [
    (
        "Hall Pass / Tardy Pass",
        "Issue and monitor digital hall passes and tardy passes in real time. A live "
        "queue shows every student currently out of class with color-coded timers "
        "(green = active, amber = nearing the limit, red = overdue), and the queue "
        "auto-resets on the bell schedule so each period starts clean.",
        "\u201cHall Pass\u201d / \u201cTardy Pass\u201d tiles on the dashboard, or the Hall Pass "
        "item in the sidebar. Admins manage the school-wide queue under Admin Hub \u2192 "
        "Hall Pass Queue.",
        "Click Create Pass \u2192 pick a student \u2192 choose a destination (e.g., Nurse) \u2192 "
        "set a max duration \u2192 Create Pass. Watch the timer change color in the live "
        "queue, then End Pass. For tardies, click Log Tardy, pick a student and reason, "
        "and print the pass.",
        [],
        "Hall Pass live queue showing active passes with color-coded timers (green / amber / red).",
    ),
    (
        "Teacher Rosters",
        "A teacher\u2019s class roster enriched with FAST performance, program flags, and "
        "grouping tools. Student pills show the full FAST journey and the tools to act "
        "on it right from the roster.",
        "Sidebar \u2192 Teacher Roster. Teachers see \u201cMy Roster\u201d; Core Team can pick "
        "any teacher and class period from the top dropdown and tabs.",
        "Pick a teacher/period \u2192 toggle Level / Scale Score \u2192 open the Benchmarks tab "
        "\u2192 Suggest Small Group \u2192 hover a program chip \u2192 Log Accommodation.",
        [
            "Accommodation logging \u2013 hover a student\u2019s ESE / 504 / ELL chip and choose "
            "Log Accommodation to record that an accommodation was used.",
            "Growth view \u2013 each pill shows Prior PM3 \u2192 PM1 \u2192 PM2 \u2192 PM3 with "
            "color-coded levels (L1\u2013L5) and a learning-gain checkmark when a student moves "
            "up a level or holds proficiency.",
            "Gap indicator \u2013 shows how many points a student needs to reach the next "
            "proficiency sub-level.",
            "Benchmark features & groupings \u2013 the Benchmarks tab groups students by mastery; "
            "Suggest Small Group auto-builds groups.",
            "Benchmark & focus-benchmark groupings \u2013 group by a chosen benchmark / focus "
            "benchmark to separate \u201cneeds support\u201d from students near a level jump.",
        ],
        "Teacher Roster with FAST pills (Prior PM3 \u2192 PM1 \u2192 PM2 \u2192 PM3) and the Benchmarks grouping view.",
    ),
    (
        "The Bell \u2013 Tier 2 & Tier 3 Interventions",
        "The bell icon at the top of every page is the MTSS reminder. It glows and "
        "shakes when a teacher has intervention logs owed for the week. Clicking it "
        "opens the interventions workspace to log Tier 2 (daily) and Tier 3 (weekly) "
        "entries; saving an entry clears the reminder.",
        "Top-right of the global header, next to the user profile, on every page.",
        "Click the bell \u2192 choose a student \u2192 complete the Tier 2 daily or Tier 3 "
        "weekly form \u2192 Save Entry, and watch the bell\u2019s count drop.",
        [],
        "The Bell icon in the header (glowing) plus the Tier 2 / Tier 3 intervention log form.",
    ),
    (
        "PBIS Points, Houses & Stores",
        "Award PBIS points to students individually or in bulk against your rubric, "
        "track the school-wide House competition, and let students redeem points in two "
        "catalogs \u2013 a School Store (school-wide) and a Classroom Store (teacher-managed "
        "rewards).",
        "Sidebar \u2192 PBIS Points. House standings also appear in the Houses panel and on "
        "signage displays.",
        "In the Classes tab, select students \u2192 Award Points \u2192 pick a rubric reason. "
        "Then open the School Store and Classroom Store tabs to browse rewards, and check "
        "the House standings.",
        [
            "House competition \u2013 live school-wide standings to drive friendly competition.",
            "Bulk awarding \u2013 select multiple students to award an entire group at once.",
            "Two stores \u2013 School Store (school-wide catalog, read-only for teachers) and "
            "Classroom Store (each teacher manages their own rewards).",
        ],
        "PBIS Points hub showing the award flow, House standings, and the School / Classroom Store tabs.",
    ),
    (
        "Academic Trajectories (Insights)",
        "A cohort view of how students moved across FAST windows, grouped into journey "
        "archetypes \u2013 Climbed, Held the Line (High), Slipped, Stuck, Held the Line "
        "(Low), and Untested \u2013 on an interactive PM1-vs-PM3 matrix. Drill into any "
        "archetype to see sub-groups and the exact student list.",
        "Sidebar \u2192 Insights \u2192 Academic Trajectory.",
        "Choose Subject (ELA / Math) and grades \u2192 click an archetype card such as "
        "Slipped or Climbed \u2192 open a sub-category \u2192 view the student list and drill "
        "into a profile.",
        [],
        "Academic Trajectory matrix with the journey archetype cards (Climbed, Slipped, Stuck, etc.).",
    ),
    (
        "Instructional Coverage (Insights)",
        "Combines teacher instruction logs with student FAST mastery to show which "
        "standards are actually being taught and how well students master them. "
        "Benchmarks are flagged in effectiveness bands (Critical, Re-teach, Building, "
        "Effective) so leaders can spot untaught or low-mastery standards at a glance.",
        "Sidebar \u2192 Insights \u2192 Instructional Coverage.",
        "Pick subject / grade \u2192 sort by \u201cWeak + untaught first\u201d \u2192 review the "
        "effectiveness bands \u2192 click a benchmark to see which teachers taught it and "
        "their mastery rates.",
        [],
        "Instructional Coverage dashboard with effectiveness bands and a benchmark drill-down.",
    ),
]

DEMO_STEPS = [
    ("Open the site (30 sec)",
     "Go to %s, scroll down, and choose PulseEDU. Point out that this is the same "
     "front door every staff member uses." % SITE),
    ("Log in as a teacher (1 min)",
     "Use an ELA login, e.g. kathleen.taylor@pulsedemo.com with the password %s. "
     "Land on the dashboard and orient the audience to the tiles and the sidebar." % PASSWORD),
    ("Hall Pass / Tardy Pass (2 min)",
     "Create a hall pass, show the timer turning amber/red in the live queue, then "
     "End Pass. Log a tardy and show the printable pass."),
    ("Teacher Roster (3 min)",
     "Show the FAST pills (Prior PM3 \u2192 PM1 \u2192 PM2 \u2192 PM3), the learning-gain "
     "checkmarks, and the gap indicator. Hover a program chip and Log Accommodation. "
     "Open the Benchmarks tab and run Suggest Small Group."),
    ("The Bell \u2013 Interventions (2 min)",
     "Click the bell in the header, show the Tier 2 / Tier 3 logs owed, and log one "
     "entry so the audience sees the count clear."),
    ("PBIS, Houses & Stores (2 min)",
     "Award points to a small group, show the House standings, then open the School "
     "Store and Classroom Store tabs."),
    ("Insights \u2013 Academic Trajectory (2 min)",
     "Go to Insights \u2192 Academic Trajectory. Pick a subject, click an archetype "
     "(e.g., Slipped), and drill into the student list."),
    ("Insights \u2013 Instructional Coverage (2 min)",
     "Go to Insights \u2192 Instructional Coverage. Sort by \u201cWeak + untaught first\u201d "
     "and drill into a benchmark to show teacher-level mastery."),
    ("Wrap (1 min)",
     "Mention that every account is scoped to a single school, so demo users can "
     "explore freely without affecting other schools or district data. To show Math "
     "data, log in again with a Math teacher, e.g. amy.brown@pulsedemo.com."),
]

ADMIN_INTRO = (
    "Short answer: give the person the viewing and logging permissions they need, but "
    "leave the three top-level admin switches OFF. PulseEDU separates what someone can "
    "see and do (capability permissions) from full administrative power (the Admin, "
    "District Admin, and SuperUser switches). The destructive tools live behind those "
    "three switches, so a capability-only account can use the app fully but cannot "
    "damage your data."
)
ADMIN_STEPS = [
    "Sign in as an admin and go to Settings \u2192 Staff & Roles.",
    "Find the staff member and open their roles.",
    "Turn ON the capability / visibility options you want them to use \u2013 for example "
    "reports, student activity, and intervention logging.",
    "Leave Admin, District Admin, and SuperUser OFF.",
    "Save. They\u2019ll see the dashboards, rosters, PBIS, interventions, and reports, and "
    "can log day-to-day data, but the Admin Hub and Data Management destructive tools "
    "stay hidden from them.",
]
ADMIN_SAFETY = [
    "What stays protected: deleting students or staff, rolling back or wiping data "
    "imports, reseeding a school, and changing core settings all require Admin, District "
    "Admin, or SuperUser \u2013 a capability-only user cannot reach them.",
    "Built-in tenant safety net: even a full school Admin can only ever affect their "
    "OWN school\u2019s data. No school-level account can touch another school or the whole "
    "district. Your Parrott demo users are sandboxed to Parrott automatically.",
    "No self-promotion: the system blocks privilege escalation \u2013 a non-admin cannot "
    "turn themselves into an admin, even if they can edit staff roles.",
    "Bottom line for demos: capability-only is the safe choice. Reserve Admin / District "
    "Admin / SuperUser for the few people you trust with deletions and imports.",
]

# ===========================================================================
# DOCX
# ===========================================================================
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_border(cell, color, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def docx_placeholder(doc, caption):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(6.4)
    shade_cell(cell, BOXBG)
    set_cell_border(cell, BOXBORDER, 8)
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("[ SCREENSHOT ]")
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = rgb(PRIMARY)
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(caption)
    r1.italic = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = rgb(MUTE)
    # add vertical breathing room inside the box
    p1.paragraph_format.space_after = Pt(18)
    p0.paragraph_format.space_before = Pt(18)
    doc.add_paragraph()

def docx_login_table(doc, rows):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Teacher", "Subject / Grade", "Username (email)"]):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        shade_cell(hdr[i], PRIMARY)
        run.font.color.rgb = rgb("FFFFFF")
    for name, subj, email in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(name).font.size = Pt(9.5)
        c[1].paragraphs[0].add_run(subj).font.size = Pt(9.5)
        er = c[2].paragraphs[0].add_run(email)
        er.font.size = Pt(9.5)
    doc.add_paragraph()

def build_docx(path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("PulseEDU")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = rgb(PRIMARY)
    sub = doc.add_paragraph()
    sr = sub.add_run("Demo Login & Feature Guide")
    sr.font.size = Pt(14)
    sr.font.color.rgb = rgb(ACCENT)
    sr.bold = True
    sch = doc.add_paragraph()
    scr = sch.add_run(SCHOOL)
    scr.font.size = Pt(10.5)
    scr.font.color.rgb = rgb(MUTE)
    doc.add_paragraph()

    # 1. Login
    doc.add_heading("1.  How to Log In", level=1)
    steps = [
        ("Go to %s" % SITE, "Open a web browser and visit the address above."),
        ("Scroll down and choose PulseEDU", "On the landing page, scroll down and select the PulseEDU product."),
        ("Pick a login below", "Use any one of the ELA or Math usernames listed on this page."),
        ("Enter the password", "The password for every demo account is:  %s" % PASSWORD),
    ]
    for i, (h, d) in enumerate(steps, 1):
        pp = doc.add_paragraph()
        rb = pp.add_run("Step %d \u2014 %s" % (i, h))
        rb.bold = True
        rb.font.color.rgb = rgb(PRIMARY)
        dd = doc.add_paragraph(d)
        dd.paragraph_format.space_after = Pt(4)

    # password callout
    ct = doc.add_table(rows=1, cols=1)
    cc = ct.cell(0, 0)
    shade_cell(cc, ACCENT)
    cpara = cc.paragraphs[0]
    cpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cpara.add_run("Password for all demo accounts:  %s" % PASSWORD)
    cr.bold = True
    cr.font.size = Pt(12)
    cr.font.color.rgb = rgb("FFFFFF")
    doc.add_paragraph()

    docx_placeholder(doc, "%s landing page \u2014 capture the page, then the spot where you scroll down to choose PulseEDU." % SITE)

    doc.add_heading("ELA Teachers (Parrott)", level=2)
    docx_login_table(doc, LOGINS_ELA)
    doc.add_heading("Math Teachers (Parrott)", level=2)
    docx_login_table(doc, LOGINS_MATH)

    # 2. Features
    doc.add_page_break()
    doc.add_heading("2.  Feature Overview", level=1)
    intro = doc.add_paragraph(
        "A quick tour of the features to show during the demo. Each section ends with a "
        "labeled box \u2014 drop in a screenshot of that screen.")
    intro.runs[0].italic = True
    for title, overview, where, try_this, bullets, cap in FEATURES:
        doc.add_heading(title, level=2)
        doc.add_paragraph(overview)
        if bullets:
            for b in bullets:
                doc.add_paragraph(b, style="List Bullet")
        wp = doc.add_paragraph()
        wp.add_run("Where to find it:  ").bold = True
        wp.add_run(where)
        tp = doc.add_paragraph()
        tp.add_run("Try this:  ").bold = True
        tp.add_run(try_this)
        docx_placeholder(doc, cap)

    # 3. Admin demo script
    doc.add_page_break()
    doc.add_heading("3.  Step-by-Step Admin Demo Script", level=1)
    doc.add_paragraph(
        "A suggested ~15-minute flow. Times are approximate \u2014 adjust to your audience.")
    for i, (h, d) in enumerate(DEMO_STEPS, 1):
        pp = doc.add_paragraph()
        rb = pp.add_run("%d.  %s" % (i, h))
        rb.bold = True
        rb.font.color.rgb = rgb(PRIMARY)
        doc.add_paragraph(d).paragraph_format.space_after = Pt(6)

    # 4. Admin access
    doc.add_page_break()
    doc.add_heading("4.  Giving Admin-Style Access Without Risking the Database", level=1)
    doc.add_paragraph(ADMIN_INTRO)
    doc.add_heading("How to set it up", level=2)
    for s in ADMIN_STEPS:
        doc.add_paragraph(s, style="List Number")
    doc.add_heading("Why this is safe", level=2)
    for s in ADMIN_SAFETY:
        doc.add_paragraph(s, style="List Bullet")

    doc.save(path)
    print("wrote", path)

# ===========================================================================
# PDF
# ===========================================================================
def build_pdf(path):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=10.5, leading=15, textColor=hx(INK), spaceAfter=6)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName="Helvetica-Bold",
                        fontSize=16, leading=20, textColor=hx(PRIMARY), spaceBefore=10, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                        fontSize=12.5, leading=16, textColor=hx(ACCENT), spaceBefore=8, spaceAfter=4)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=2, spaceAfter=3)
    boxcap = ParagraphStyle("boxcap", parent=body, alignment=TA_CENTER, fontSize=9,
                            textColor=hx(MUTE), leading=12)
    boxlbl = ParagraphStyle("boxlbl", parent=body, alignment=TA_CENTER, fontSize=10,
                            textColor=hx(PRIMARY), fontName="Helvetica-Bold", spaceAfter=4)
    strong = ParagraphStyle("strong", parent=body, fontName="Helvetica-Bold")

    story = []

    def placeholder(cap):
        inner = [Paragraph("[ SCREENSHOT ]", boxlbl), Paragraph(cap, boxcap)]
        t = Table([[inner]], colWidths=[6.6 * inch], rowHeights=[1.15 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), hx(BOXBG)),
            ("BOX", (0, 0), (-1, -1), 1, hx(BOXBORDER)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ]))
        return t

    # Title block
    story.append(Paragraph("PulseEDU", ParagraphStyle(
        "title", parent=h1, fontSize=28, leading=32, spaceAfter=2)))
    story.append(Paragraph("Demo Login &amp; Feature Guide", ParagraphStyle(
        "subtitle", parent=h2, fontSize=15, spaceBefore=0, spaceAfter=2)))
    story.append(Paragraph(SCHOOL, ParagraphStyle(
        "school", parent=body, textColor=hx(MUTE))))
    story.append(Spacer(1, 14))

    # 1. Login
    story.append(Paragraph("1.&nbsp;&nbsp;How to Log In", h1))
    login_steps = [
        ("Go to %s" % SITE, "Open a web browser and visit the address above."),
        ("Scroll down and choose PulseEDU", "On the landing page, scroll down and select the PulseEDU product."),
        ("Pick a login below", "Use any one of the ELA or Math usernames listed on this page."),
        ("Enter the password", "The password for every demo account is shown below."),
    ]
    for i, (h, d) in enumerate(login_steps, 1):
        story.append(Paragraph("<font color='#%s'><b>Step %d &mdash; %s</b></font>" % (PRIMARY, i, h), body))
        story.append(Paragraph(d, body))

    # password callout
    pw = Table([[Paragraph(
        "<font color='white'><b>Password for all demo accounts:&nbsp;&nbsp;%s</b></font>" % PASSWORD,
        ParagraphStyle("pw", parent=body, alignment=TA_CENTER, fontSize=13))]],
        colWidths=[6.6 * inch])
    pw.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), hx(ACCENT)),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(Spacer(1, 4))
    story.append(pw)
    story.append(Spacer(1, 10))
    story.append(placeholder(
        "%s landing page \u2014 capture the page, then the spot where you scroll down to choose PulseEDU." % SITE))
    story.append(Spacer(1, 12))

    def login_table(rows, heading):
        story.append(Paragraph(heading, h2))
        data = [[Paragraph("<b>Teacher</b>", ParagraphStyle("th", parent=body, textColor=colors.white, fontSize=9.5)),
                 Paragraph("<b>Subject / Grade</b>", ParagraphStyle("th2", parent=body, textColor=colors.white, fontSize=9.5)),
                 Paragraph("<b>Username (email)</b>", ParagraphStyle("th3", parent=body, textColor=colors.white, fontSize=9.5))]]
        cellst = ParagraphStyle("cell", parent=body, fontSize=9.5, spaceAfter=0)
        for name, subj, email in rows:
            data.append([Paragraph(name, cellst), Paragraph(subj, cellst), Paragraph(email, cellst)])
        t = Table(data, colWidths=[1.8 * inch, 1.6 * inch, 3.2 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), hx(PRIMARY)),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, hx("EEF1F9")]),
            ("GRID", (0, 0), (-1, -1), 0.5, hx(BOXBORDER)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    login_table(LOGINS_ELA, "ELA Teachers (Parrott)")
    login_table(LOGINS_MATH, "Math Teachers (Parrott)")

    # 2. Features
    story.append(Paragraph("2.&nbsp;&nbsp;Feature Overview", h1))
    story.append(Paragraph(
        "<i>A quick tour of the features to show during the demo. Each section ends with a "
        "labeled box &mdash; drop in a screenshot of that screen.</i>", body))
    for title, overview, where, try_this, bullets, cap in FEATURES:
        block = [Paragraph(title, h2), Paragraph(overview, body)]
        for b in bullets:
            block.append(Paragraph("&bull;&nbsp;&nbsp;" + b, bullet))
        block.append(Paragraph("<b>Where to find it:</b>&nbsp;&nbsp;" + where, body))
        block.append(Paragraph("<b>Try this:</b>&nbsp;&nbsp;" + try_this, body))
        block.append(Spacer(1, 4))
        block.append(placeholder(cap))
        block.append(Spacer(1, 14))
        story.append(KeepTogether(block))

    # 3. Demo script
    story.append(Paragraph("3.&nbsp;&nbsp;Step-by-Step Admin Demo Script", h1))
    story.append(Paragraph(
        "A suggested ~15-minute flow. Times are approximate &mdash; adjust to your audience.", body))
    for i, (h, d) in enumerate(DEMO_STEPS, 1):
        story.append(Paragraph("<font color='#%s'><b>%d.&nbsp;&nbsp;%s</b></font>" % (PRIMARY, i, h), body))
        story.append(Paragraph(d, body))
        story.append(Spacer(1, 2))

    # 4. Admin access
    story.append(Paragraph("4.&nbsp;&nbsp;Giving Admin-Style Access Without Risking the Database", h1))
    story.append(Paragraph(ADMIN_INTRO, body))
    story.append(Paragraph("How to set it up", h2))
    for i, s in enumerate(ADMIN_STEPS, 1):
        story.append(Paragraph("<b>%d.</b>&nbsp;&nbsp;%s" % (i, s), bullet))
    story.append(Paragraph("Why this is safe", h2))
    for s in ADMIN_SAFETY:
        story.append(Paragraph("&bull;&nbsp;&nbsp;" + s, bullet))

    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=0.8 * inch, rightMargin=0.8 * inch,
                            topMargin=0.7 * inch, bottomMargin=0.7 * inch,
                            title="PulseEDU Demo Login & Feature Guide")
    doc.build(story)
    print("wrote", path)


if __name__ == "__main__":
    build_docx("exports/PulseEDU-Demo-Guide.docx")
    build_pdf("exports/PulseEDU-Demo-Guide.pdf")
